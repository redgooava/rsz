import tkinter
from tkinter.ttk import Combobox

import psycopg2
from tkinter import *
from tkinter import ttk
from tkcalendar import Calendar, DateEntry
import datetime

from docx import Document
from docx.shared import Mm
from docx.enum.section import WD_ORIENTATION, WD_SECTION_START
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from docx.shared import Pt, RGBColor, Mm, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


con = psycopg2.connect(
    database="postgres",
    user="postgres",
    password="",
    host="127.0.0.1",
    port="5432"
)
cur = con.cursor()

# окошко
window = Tk()
window.geometry('1000x750')
window.title("Строевая записка")

# менюбар
tab_control = ttk.Notebook(window)
tab1 = ttk.Frame(tab_control)
tab3 = ttk.Frame(tab_control)
tab7 = ttk.Frame(tab_control)
tab8 = ttk.Frame(tab_control)
tab9 = ttk.Frame(tab_control)
tab10 = ttk.Frame(tab_control)
tab11 = ttk.Frame(tab_control)
tab_control.add(tab1, text='Таблица')
tab_control.add(tab3, text='Отсутствие')
tab_control.add(tab7, text='Поиск человека')
tab_control.add(tab8, text='Поиск подразделения')
tab_control.add(tab9, text='Удаление подразделения')
tab_control.add(tab10, text='Добавление подразделения')
tab_control.add(tab11, text='О программе')
tab_control.pack(expand=1, fill='both')

actionRows = []
actionRows2 = []

listOfRanks = ('Оф', 'Пр', 'К/с', 'С/с', 'К-ты', 'Сл')

listOfDivisions = []

listOfReasons = ['Отпуск', 'Командировка', 'Амбулаторное лечение', 'Стационарное лечение', 'Наряд']


def initOfDivisions(listOfDivisions):
    listOfDivisions.clear()

    cur.execute(
        f"SELECT * FROM DIVISIONTABLE"
    )
    con.commit()

    listOfDivisionsPre = list(cur.fetchall())

    for i in listOfDivisionsPre:
        listOfDivisions.append(i[1])


# добавление записи (для отпуска, командировки и прочих записей об отсутствии)
def actionInput(name, whyout, start, end, sDivision, numberOfOrder, rank):
    if (name is not None) and (whyout is not None) and (start is not None) and (end is not None):
        cur.execute(
            f"INSERT INTO OUTTABLE (NAME, WHYOUT, DATESTART, DATEEND, SUBDIVISION, NUMBEROFORDER, RANK) "
            f"VALUES ('{name}', '{whyout}', '{start}', '{end}', '{sDivision}', '{numberOfOrder}', '{rank}')"
        )
        con.commit()
    else:
        print('Ошибка в actionInput')
        
        play()


def actionSearch(name, table):
    cur.execute(
        f"SELECT * FROM OUTTABLE "
        f"WHERE name = '{name}'"
    )
    for item in table.get_children():
        table.delete(item)

    actionRows = cur.fetchall()
    con.commit()

    for row in actionRows:
        table.insert('', tkinter.END, values=row)

    table.pack(expand=tkinter.YES, fill=tkinter.BOTH)
    table.place(x=0, y=200)

    window.update()


def actionSearchDivision(division, table):
    cur.execute(
        f"SELECT * FROM OUTTABLE "
        f"WHERE subdivision = '{division}'"
    )

    for item in table.get_children():
        table.delete(item)

    actionRows = cur.fetchall()
    con.commit()

    for row in actionRows:
        table.insert('', tkinter.END, values=row)

    table.pack(expand=tkinter.YES, fill=tkinter.BOTH)
    table.place(x=0, y=200)


def actionEdit(division, officer, ensign, contract, soldier, cadet, listener, total,
               s_officer, s_ensign, s_contract, s_soldier, s_cadet, s_listener, s_total):
    cur.execute(
        f"INSERT INTO DIVISIONTABLE (DIVISION, OFFICER, ENSIGN, CONTRACT, SOLDIER, CADET, LISTENER, TOTAL, "
        f" SOFFICER, SENSIGN, SCONTRACT, SSOLDIER, SCADET, SLISTENER, STOTAL)"
        f"VALUES('{division}', '{officer}', '{ensign}', '{contract}', '{soldier}', '{cadet}', '{listener}', '{total}', "
        f"'{s_officer}', '{s_ensign}', '{s_contract}', '{s_soldier}', '{s_cadet}', '{s_listener}', '{s_total}')"
    )
    con.commit()
    initOfDivisions(listOfDivisions)

    play()


def actionGeneral(table):
    for item in table.get_children():
        table.delete(item)

    cur.execute(
        f"SELECT * FROM DIVISIONTABLE"
    )
    con.commit()

    actionRows = cur.fetchall()

    for i in range(len(actionRows)):
        actionRows[i] = list(actionRows[i])
        actionRows[i].append(actionRows[i][9] - inPeriodAtDivision(actionRows[i][1], 'Оф'))
        actionRows[i].append(actionRows[i][10] - inPeriodAtDivision(actionRows[i][1], 'Пр'))
        actionRows[i].append(actionRows[i][11] - inPeriodAtDivision(actionRows[i][1], 'К/с'))
        actionRows[i].append(actionRows[i][12] - inPeriodAtDivision(actionRows[i][1], 'С/с'))
        actionRows[i].append(actionRows[i][13] - inPeriodAtDivision(actionRows[i][1], 'К-ты'))
        actionRows[i].append(actionRows[i][14] - inPeriodAtDivision(actionRows[i][1], 'Сл'))
        actionRows[i].append(actionRows[i][16] + actionRows[i][17] + actionRows[i][18] +
                             actionRows[i][19] + actionRows[i][20] + actionRows[i][21])
        actionRows[i].append(inPeriodNotAtDivision(actionRows[i][1], 'Отпуск'))
        actionRows[i].append(inPeriodNotAtDivision(actionRows[i][1], 'Стационарное лечение'))
        actionRows[i].append(inPeriodNotAtDivision(actionRows[i][1], 'Амбулаторное лечение'))
        actionRows[i].append(inPeriodNotAtDivision(actionRows[i][1], 'Наряд'))
        actionRows[i].append(inPeriodNotAtDivision(actionRows[i][1], 'Командировка'))
        actionRows[i].append(actionRows[i][23] + actionRows[i][24] +
                             actionRows[i][25] + actionRows[i][26] + actionRows[i][27])

    for row in actionRows:
        table.insert('', tkinter.END, values=row)

    table.pack(expand=tkinter.YES, fill=tkinter.BOTH)
    table.place(x=40, y=50, width=900, height=400)


def inPeriodAtDivision(division, rank):
    print(division + ' ' + rank)
    cur.execute(
        f"SELECT * FROM OUTTABLE "
        f"WHERE subdivision = '{division}' AND (rank = '{rank}') AND "
        f"(datestart <= now()) AND "
        f"(dateend > now())"
    )
    con.commit()

    actionRows = cur.fetchall()
    return len(actionRows)


def inPeriodNotAtDivision(division, whyout):
    cur.execute(
        f"SELECT * FROM OUTTABLE "
        f"WHERE subdivision = '{division}' AND whyout = '{whyout}' AND "
        f"(datestart <= now()) AND "
        f"(dateend > now())"
    )
    con.commit()

    actionRows = cur.fetchall()
    return len(actionRows)


def deleteDivision(division):
    cur.execute(
        f"DELETE FROM DIVISIONTABLE "
        f"WHERE division = '{division}'"
    )
    con.commit()

    play()


def toDocx():  # https://docs-python.ru/packages/modul-python-docx-python/izmenenie-maketa-dokumenta/
    # doc = Document()
    # style = doc.styles['Normal']
    # style.font.size = Pt(8)
    # section = doc.sections[0]
    # # левое поле в миллиметрах
    # section.left_margin = Mm(10)
    # # правое поле в миллиметрах
    # section.right_margin = Mm(10)
    # # верхнее поле в миллиметрах
    # section.top_margin = Mm(20)
    # # нижнее поле в миллиметрах
    # section.bottom_margin = Mm(20)
    # doc.add_section(start_type=WD_SECTION_START.NEW_PAGE)
    # section.orientation = WD_ORIENTATION.LANDSCAPE
    # doc.sections[1].page_width = doc.sections[0].page_height
    # doc.sections[1].page_height = doc.sections[0].page_width
    # head = doc.add_heading('Развёрнутая строевая записка')
    # head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    #
    # pretable = doc.add_table(rows=1, cols=5)
    # pretable.rows[0].cells[0].text = ''
    # pretable.rows[0].cells[1].text = 'По штату'
    # pretable.rows[0].cells[2].text = 'По списку'
    # pretable.rows[0].cells[3].text = 'Налицо'
    # pretable.rows[0].cells[4].text = 'Отсутствуют'
    #
    # table = doc.add_table(rows=len(divisionlist) + 1, cols=28)
    # table.autofit = False
    # table.rows[0].cells[0].text = 'Подразделение'
    # # table.rows[0].cells[1].text = 'По штату'
    # table.rows[0].cells[1].text = 'Оф'
    # table.rows[0].cells[2].text = 'Пр'
    # table.rows[0].cells[3].text = 'К/с'
    # table.rows[0].cells[4].text = 'С/с'
    # table.rows[0].cells[5].text = 'К-ты'
    # table.rows[0].cells[6].text = 'Сл'
    # table.rows[0].cells[7].text = 'ВСЕГО'
    # # table.rows[0].cells[8].text = 'По списку'
    # table.rows[0].cells[8].text = 'Оф'
    # table.rows[0].cells[9].text = 'Пр'
    # table.rows[0].cells[10].text = 'К/с'
    # table.rows[0].cells[11].text = 'С/с'
    # table.rows[0].cells[12].text = 'К-ты'
    # table.rows[0].cells[13].text = 'Сл'
    # table.rows[0].cells[14].text = 'ВСЕГО'
    # # table.rows[0].cells[15].text = 'Налицо'
    # table.rows[0].cells[15].text = 'Оф'
    # table.rows[0].cells[16].text = 'Пр'
    # table.rows[0].cells[17].text = 'К/с'
    # table.rows[0].cells[18].text = 'С/с'
    # table.rows[0].cells[19].text = 'К-ты'
    # table.rows[0].cells[20].text = 'Сл'
    # table.rows[0].cells[21].text = 'ВСЕГО'
    # # table.rows[0].cells[22].text = 'Отсутствуют'
    # table.rows[0].cells[22].text = 'Отпуск'
    # table.rows[0].cells[23].text = 'Стац'
    # table.rows[0].cells[24].text = 'Амб'
    # table.rows[0].cells[25].text = 'Наряд'
    # table.rows[0].cells[26].text = 'Ком-ка'
    # table.rows[0].cells[27].text = 'ВСЕГО'
    # for i in range(len(table.rows) - 1):
    #     for j in range(28 - 1):
    #         table.rows[i + 1].cells[j + 1].text = str(i + j)
    # doc.save('result.docx')

    cur.execute(
        f"SELECT * FROM DIVISIONTABLE"
    )
    con.commit()

    actionRows = cur.fetchall()

    for i in range(len(actionRows)):
        actionRows[i] = list(actionRows[i])
        actionRows[i].append(actionRows[i][9] - inPeriodAtDivision(actionRows[i][1], 'Оф'))
        actionRows[i].append(actionRows[i][10] - inPeriodAtDivision(actionRows[i][1], 'Пр'))
        actionRows[i].append(actionRows[i][11] - inPeriodAtDivision(actionRows[i][1], 'К/с'))
        print(str(actionRows[i][18]) + ' ' + str(actionRows[i][11]) + ' ' + str(inPeriodAtDivision(actionRows[i][1], 'К/с')))
        actionRows[i].append(actionRows[i][12] - inPeriodAtDivision(actionRows[i][1], 'С/с'))
        actionRows[i].append(actionRows[i][13] - inPeriodAtDivision(actionRows[i][1], 'К-ты'))
        actionRows[i].append(actionRows[i][14] - inPeriodAtDivision(actionRows[i][1], 'Сл'))
        actionRows[i].append(actionRows[i][16] + actionRows[i][17] + actionRows[i][18] +
                             actionRows[i][19] + actionRows[i][20] + actionRows[i][21])
        actionRows[i].append(inPeriodNotAtDivision(actionRows[i][1], 'Отпуск'))
        actionRows[i].append(inPeriodNotAtDivision(actionRows[i][1], 'Стационарное лечение'))
        actionRows[i].append(inPeriodNotAtDivision(actionRows[i][1], 'Амбулаторное лечение'))
        actionRows[i].append(inPeriodNotAtDivision(actionRows[i][1], 'Наряд'))
        actionRows[i].append(inPeriodNotAtDivision(actionRows[i][1], 'Командировка'))
        actionRows[i].append(actionRows[i][23] + actionRows[i][24] +
                             actionRows[i][25] + actionRows[i][26] + actionRows[i][27])

    templist = ['', 'ВСЕГО']
    for j in range(len(actionRows[0]) - 2):
        temp = 0
        for i in actionRows:
            temp += i[j + 2]
        templist.append(temp)
    actionRows.append(templist)

    doc = Document('template.docx')
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(10)

    table = doc.tables[0]

    for i in range(len(actionRows)):
        for j in range(len(actionRows[0]) - 1):
            table.cell(i + 3, j).text = str(actionRows[i][j + 1])

    doc.add_paragraph('ОТСУТСТВУЮТ: ')

    tableout = [None] * len(listOfDivisions)
    for i in range(len(listOfDivisions)):
        cur.execute(
            f"SELECT * FROM OUTTABLE "
            f"WHERE subdivision = '{listOfDivisions[i]}' AND "
            f"(datestart <= now()) AND "
            f"(dateend > now())"
        )
        con.commit()

        doc.add_paragraph(listOfDivisions[i])
        doc.add_paragraph('')

        actionRows2 = cur.fetchall()

        tableout[i] = doc.add_table(rows=len(actionRows2) + 1, cols=7)
        tableout[i].cell(0, 0).text = '№ п/п'
        tableout[i].cell(0, 1).text = 'ФИО'
        tableout[i].cell(0, 2).text = 'Причина'
        tableout[i].cell(0, 3).text = 'С какого дня'
        tableout[i].cell(0, 4).text = 'По какой день'
        tableout[i].cell(0, 5).text = 'Номер приказа'
        tableout[i].cell(0, 6).text = 'Категория'

        for j in range(len(actionRows2)):
            actionRows2[j] = list(actionRows2[j])
            actionRows2[j][4] = actionRows2[j][3]
            actionRows2[j][3] = actionRows2[j][2]
            actionRows2[j][2] = actionRows2[j][1]
            actionRows2[j][1] = actionRows2[j][0]
            actionRows2[j][0] = j + 1

        # print(len(actionRows2[0]))
        for k in range(len(actionRows2)):
            for m in range(len(actionRows2[0])):
                tableout[i].cell(k + 1, m).text = str(actionRows2[k][m])

        doc.add_page_break()

    doc.save('result.docx')


# вкладка "Отсутствие"
def trip():
    general_lbl = Label(tab3, text='Командировка', font=("Arial", 20))
    general_lbl.place(width=250, height=30, x=130, y=70)
    name_trip_lbl = Label(tab3, text="ФИО: ")
    name_trip_lbl.place(width=250, height=30, x=0, y=150)
    name_trip_ent = Entry(tab3, width=10, font=("Verdana", 12))
    name_trip_ent.place(width=250, height=30, x=180, y=150)
    reason_lbl = Label(tab3, text="Причина: ")
    reason_lbl.place(width=250, height=30, x=0, y=200)
    reason_cb = Combobox(tab3, values=listOfReasons)
    reason_cb.place(width=250, height=30, x=180, y=200)
    cal_start_trip_lbl = Label(tab3, text="Дата убытия: ")
    cal_start_trip_lbl.place(width=250, height=30, x=0, y=250)
    cal_start_trip = DateEntry(tab3, width=16, background="magenta3", foreground="white", bd=2)
    cal_start_trip.pack(pady=20)
    cal_start_trip.place(width=250, height=30, x=180, y=250)
    cal_end_trip_lbl = Label(tab3, text="Дата прибытия: ")
    cal_end_trip_lbl.place(width=250, height=30, x=0, y=300)
    cal_end_trip = DateEntry(tab3, width=16, background="magenta3", foreground="white", bd=2)
    cal_end_trip.pack(pady=20)
    cal_end_trip.place(width=250, height=30, x=180, y=300)
    division_lbl = Label(tab3, text="Подразделение: ")
    division_lbl.place(width=250, height=30, x=0, y=350)
    division = Combobox(tab3, values=listOfDivisions)
    division.place(width=250, height=30, x=180, y=350)
    numberOfOrder_lbl = Label(tab3, text="Номер приказа: ")
    numberOfOrder_lbl.place(width=250, height=30, x=0, y=400)
    numberOfOrder = Entry(tab3, width=10, font=("Verdana", 12))
    numberOfOrder.place(width=250, height=30, x=180, y=400)
    rank_lbl = Label(tab3, text="Категория: ")
    rank_lbl.place(width=250, height=30, x=0, y=450)
    rank = Combobox(tab3, values=listOfRanks)
    rank.place(width=250, height=30, x=180, y=450)
    button_trip = Button(tab3, text='Записать', command=lambda: actionInput(
        name_trip_ent.get(), reason_cb.get(),
        cal_start_trip.get_date().strftime("%d/%m/%Y"), cal_end_trip.get_date().strftime("%d/%m/%Y"), division.get(),
        numberOfOrder.get(), rank.get()))
    button_trip.place(width=100, height=30, x=330, y=500)


# вкладка "Поиск человека"
def search():
    name_search_lbl = Label(tab7, text="ФИО: ")
    name_search_lbl.place(width=250, height=30, x=0, y=50)
    name_search_ent = Entry(tab7, width=10, font=("Verdana", 12))
    name_search_ent.place(width=250, height=30, x=180, y=50)

    # всё про таблицу отсюда https://www.youtube.com/watch?v=HMPIeZ3S_cs
    table = ttk.Treeview(tab7, show='headings')
    table.place(x=0, y=200)
    heads = ['ФИО', 'Причина отсутствия', 'C какого дня', 'По какой день',
             'Подразделение', 'Номера приказа', 'Категория']
    table['columns'] = heads

    for header in heads:
        table.heading(header, text=header, anchor='center')
        table.column(header, anchor='center', width=130)

    scroll_pane = ttk.Scrollbar(tab7, command=table.yview)
    table.configure(yscrollcommand=scroll_pane.set)
    scroll_pane.pack(side=tkinter.RIGHT, fill=tkinter.Y)
    scroll_pane.place(height=220, x=1204, y=200)

    button_search = Button(tab7, text='Найти', command=lambda: actionSearch(name_search_ent.get(), table))
    button_search.place(width=100, height=30, x=450, y=50)


def searchDivision():
    division_search_lbl = Label(tab8, text="Подразделение: ")
    division_search_lbl.place(width=250, height=30, x=0, y=50)
    division_search_cb = Combobox(tab8, values=listOfDivisions)
    division_search_cb.place(width=250, height=30, x=180, y=50)

    # всё про таблицу отсюда https://www.youtube.com/watch?v=HMPIeZ3S_cs
    dtable = ttk.Treeview(tab8, show='headings')
    dtable.place(x=0, y=200)
    heads = ['ФИО', 'Причина отсутствия', 'C какого дня', 'По какой день',
             'Подразделение', 'Номера приказа', 'Категория']
    dtable['columns'] = heads

    for header in heads:
        dtable.heading(header, text=header, anchor='center')
        dtable.column(header, anchor='center', width=130)

    scroll_pane = ttk.Scrollbar(tab8, command=dtable.yview)
    dtable.configure(yscrollcommand=scroll_pane.set)
    scroll_pane.pack(side=tkinter.RIGHT, fill=tkinter.Y)
    scroll_pane.place(height=220, x=1204, y=200)

    button_search = Button(tab8, text='Найти', command=lambda: actionSearchDivision(division_search_cb.get(), dtable))
    button_search.place(width=100, height=30, x=450, y=50)


def deleteFromDivisionTable():
    initOfDivisions(listOfDivisions)
    general_lbl = Label(tab9, text='Удаление подразделения', font=("Arial", 20))
    general_lbl.place(width=380, height=30, x=100, y=70)

    division_lbl = Label(tab9, text="Подразделение: ")
    division_lbl.place(width=250, height=30, x=0, y=150)
    division_cb = Combobox(tab9, values=listOfDivisions)
    division_cb.place(width=250, height=30, x=220, y=150)

    button_division = Button(tab9, text='Удалить', command=lambda: deleteDivision(division_cb.get()))
    button_division.place(width=100, height=30, x=490, y=150)

    initOfDivisions(listOfDivisions)


def generalTable():
    columns = ("#1", "#2", "#3", "#4", "#5", "#6", "#7", "#8", "#9", "#10", "#11", "#12", "#13", "#14", "#15", "#16",
               "#17", "#18", "#19", "#20", "#21", "#22", "#23", "#24", "#25", "#26", "#27", "#28", "#29")
    table = ttk.Treeview(tab1, show="headings", columns=columns)
    table.place(x=40, y=50, width=900, height=400)
    table.heading("#1", text="№п/п")
    table.column("#1", anchor='center', width=40)
    table.heading("#2", text="Подразделение")
    table.column("#2", anchor='center', width=120)
    table.heading("#3", text="Оф")
    table.column("#3", anchor='center', width=50)
    table.heading("#4", text="Пр")
    table.column("#4", anchor='center', width=50)
    table.heading("#5", text="К/с")
    table.column("#5", anchor='center', width=50)
    table.heading("#6", text="С/с")
    table.column("#6", anchor='center', width=50)
    table.heading("#7", text="К-ты")
    table.column("#7", anchor='center', width=50)
    table.heading("#8", text="Сл")
    table.column("#8", anchor='center', width=50)
    table.heading("#9", text="ВСЕГО")
    table.column("#9", anchor='center', width=50)
    table.heading("#10", text="Оф")
    table.column("#10", anchor='center', width=50)
    table.heading("#11", text="Пр")
    table.column("#11", anchor='center', width=50)
    table.heading("#12", text="К/с")
    table.column("#12", anchor='center', width=50)
    table.heading("#13", text="С/с")
    table.column("#13", anchor='center', width=50)
    table.heading("#14", text="К-ты")
    table.column("#14", anchor='center', width=50)
    table.heading("#15", text="Сл")
    table.column("#15", anchor='center', width=50)
    table.heading("#16", text="ВСЕГО")
    table.column("#16", anchor='center', width=50)
    table.heading("#17", text="Оф")
    table.column("#17", anchor='center', width=50)
    table.heading("#18", text="Пр")
    table.column("#18", anchor='center', width=50)
    table.heading("#19", text="К/с")
    table.column("#19", anchor='center', width=50)
    table.heading("#20", text="С/с")
    table.column("#20", anchor='center', width=50)
    table.heading("#21", text="К-ты")
    table.column("#21", anchor='center', width=50)
    table.heading("#22", text="Сл")
    table.column("#22", anchor='center', width=50)
    table.heading("#23", text="ВСЕГО")
    table.column("#23", anchor='center', width=50)
    table.heading("#24", text="Отпуск")
    table.column("#24", anchor='center', width=50)
    table.heading("#25", text="Стац")
    table.column("#25", anchor='center', width=50)
    table.heading("#26", text="Амб")
    table.column("#26", anchor='center', width=50)
    table.heading("#27", text="Наряд")
    table.column("#27", anchor='center', width=50)
    table.heading("#28", text="Ком-ка")
    table.column("#28", anchor='center', width=50)
    table.heading("#29", text="ВСЕГО")
    table.column("#29", anchor='center', width=50)

    scroll_pane = ttk.Scrollbar(tab1, command=table.yview)
    table.configure(yscrollcommand=scroll_pane.set)
    scroll_pane.pack(side=tkinter.RIGHT, fill=tkinter.Y)
    scroll_pane.place(height=400, x=940, y=50)

    scroll_pane1 = ttk.Scrollbar(tab1, command=table.xview, orient='horizontal')
    table.configure(xscrollcommand=scroll_pane1.set)
    scroll_pane1.pack(side=tkinter.BOTTOM, fill=tkinter.X)
    scroll_pane1.place(width=900, x=40, y=450)

    button_general = Button(tab1, text='Сформировать', command=lambda: actionGeneral(table))
    button_general.place(width=100, height=30, x=820, y=490)

    button_print = Button(tab1, text='Экспорт', command=lambda: toDocx())
    button_print.place(width=100, height=30, x=820, y=540)


def editDivisionTable():
    general_lbl = Label(tab10, text='Добавление подразделения', font=("Arial", 20))
    general_lbl.place(width=380, height=30, x=100, y=70)

    division_lbl = Label(tab10, text="Подразделение: ")
    division_lbl.place(width=250, height=30, x=0, y=150)
    # division_cb = Combobox(tab10, values=listOfDivisions)
    division_ent = Entry(tab10, width=10, font=("Verdana", 12))
    division_ent.place(width=250, height=30, x=220, y=150)

    officer_lbl = Label(tab10, text="Количество офицеров: ")
    officer_lbl.place(width=250, height=30, x=0, y=200)
    officer_ent = Entry(tab10, width=10, font=("Verdana", 12))
    officer_ent.place(width=250, height=30, x=220, y=200)

    ensign_lbl = Label(tab10, text="Количество прапорщиков: ")
    ensign_lbl.place(width=250, height=30, x=0, y=250)
    ensign_ent = Entry(tab10, width=10, font=("Verdana", 12))
    ensign_ent.place(width=250, height=30, x=220, y=250)

    contract_lbl = Label(tab10, text="Количество в/с к/с: ")
    contract_lbl.place(width=250, height=30, x=0, y=300)
    contract_ent = Entry(tab10, width=10, font=("Verdana", 12))
    contract_ent.place(width=250, height=30, x=220, y=300)

    soldier_lbl = Label(tab10, text="Количество  в/с с/с: ")
    soldier_lbl.place(width=250, height=30, x=0, y=350)
    soldier_ent = Entry(tab10, width=10, font=("Verdana", 12))
    soldier_ent.place(width=250, height=30, x=220, y=350)

    cadet_lbl = Label(tab10, text="Количество курсантов: ")
    cadet_lbl.place(width=250, height=30, x=0, y=400)
    cadet_ent = Entry(tab10, width=10, font=("Verdana", 12))
    cadet_ent.place(width=250, height=30, x=220, y=400)

    listener_lbl = Label(tab10, text="Количество слушателей: ")
    listener_lbl.place(width=250, height=30, x=0, y=450)
    listener_ent = Entry(tab10, width=10, font=("Verdana", 12))
    listener_ent.place(width=250, height=30, x=220, y=450)

    total_lbl = Label(tab10, text="Общее количество: ")
    total_lbl.place(width=250, height=30, x=0, y=500)
    total_ent = Entry(tab10, width=10, font=("Verdana", 12))
    total_ent.place(width=250, height=30, x=220, y=500)

    s_officer_lbl = Label(tab10, text="Количество офицеров: ")
    s_officer_lbl.place(width=250, height=30, x=400, y=200)
    s_officer_ent = Entry(tab10, width=10, font=("Verdana", 12))
    s_officer_ent.place(width=250, height=30, x=620, y=200)

    s_ensign_lbl = Label(tab10, text="Количество прапорщиков: ")
    s_ensign_lbl.place(width=250, height=30, x=400, y=250)
    s_ensign_ent = Entry(tab10, width=10, font=("Verdana", 12))
    s_ensign_ent.place(width=250, height=30, x=620, y=250)

    s_contract_lbl = Label(tab10, text="Количество в/с к/с: ")
    s_contract_lbl.place(width=250, height=30, x=400, y=300)
    s_contract_ent = Entry(tab10, width=10, font=("Verdana", 12))
    s_contract_ent.place(width=250, height=30, x=620, y=300)

    s_soldier_lbl = Label(tab10, text="Количество  в/с с/с: ")
    s_soldier_lbl.place(width=250, height=30, x=400, y=350)
    s_soldier_ent = Entry(tab10, width=10, font=("Verdana", 12))
    s_soldier_ent.place(width=250, height=30, x=620, y=350)

    s_cadet_lbl = Label(tab10, text="Количество курсантов: ")
    s_cadet_lbl.place(width=250, height=30, x=400, y=400)
    s_cadet_ent = Entry(tab10, width=10, font=("Verdana", 12))
    s_cadet_ent.place(width=250, height=30, x=620, y=400)

    s_listener_lbl = Label(tab10, text="Количество слушателей: ")
    s_listener_lbl.place(width=250, height=30, x=400, y=450)
    s_listener_ent = Entry(tab10, width=10, font=("Verdana", 12))
    s_listener_ent.place(width=250, height=30, x=620, y=450)

    s_total_lbl = Label(tab10, text="Общее количество: ")
    s_total_lbl.place(width=250, height=30, x=400, y=500)
    s_total_ent = Entry(tab10, width=10, font=("Verdana", 12))
    s_total_ent.place(width=250, height=30, x=620, y=500)

    button_division = Button(tab10, text='Добавить', command=lambda: actionEdit(
        division_ent.get(), officer_ent.get(), ensign_ent.get(), contract_ent.get(),
        soldier_ent.get(), cadet_ent.get(), listener_ent.get(), total_ent.get(),
        s_officer_ent.get(), s_ensign_ent.get(), s_contract_ent.get(),
        s_soldier_ent.get(), s_cadet_ent.get(), s_listener_ent.get(), s_total_ent.get()
    ))
    button_division.place(width=100, height=30, x=330, y=550)


def play():
    initOfDivisions(listOfDivisions)
    trip()
    search()
    searchDivision()
    deleteFromDivisionTable()
    generalTable()
    editDivisionTable()


play()
window.mainloop()
